"""Utility to read .docx files and print their text to stdout.

Usage:
    python read_docx.py path/to/document.docx

If `python-docx` is not installed the script will attempt to install it via pip.
"""
import sys
import subprocess
from typing import TYPE_CHECKING

# Help static analyzers / language servers: show import in type-checking mode
if TYPE_CHECKING:
    import docx  # pragma: no cover

try:
    import docx
except ImportError:
    # Install python-docx at runtime if missing (useful in notebooks/environments)
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    import docx

from docx import Document


def read_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # include simple table extraction
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text for cell in row.cells)
            parts.append(row_text)
    return "\n".join(p for p in parts if p)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python read_docx.py path/to/file.docx')
        sys.exit(1)
    path = sys.argv[1]
    try:
        text = read_docx(path)
        print(text)
    except Exception as e:
        print(f'Error reading {path}: {e}', file=sys.stderr)
        sys.exit(2)
